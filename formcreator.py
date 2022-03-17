"""
                 Carlos Andres Mejias Mieses
    Propositos: 1.Popular formas de word con informacion de una tabla
                de excel.
                2.Crear pdf de esas formas.

"""


# Importacion de Modulos.
# -----------------------------------------------------------------------------
import os
from datetime import date
from docx import Document
from openpyxl import load_workbook
from docx2pdf import convert
# -----------------------------------------------------------------------------


# -----------------------------------------------------------------------------
# Funcion para convertir los idomas a espanol.
def get_date():

    # Dictionary for corresponding date name from english to spanish
    spanish = {"January":"Enero", "February":"Febrero", "March":"Marzo", "April":"Abril",
                "May":"Mayo", "June":"Junio", "July":"Julio", "August":"Agosto", "September":"Septiembre",
                "October":"Octubre", "November":"Noviembre", "December":"Diciembre"}

    # guarda el dia.
    today = date.today()
    # convierte el dia a un formato manipulable.
    today = today.strftime("%d,%B,%Y")
    day, month, year = str(today).split(",")
    # Convierte los meses de ingles a esapanol.
    month = spanish[month]

    # devuelve dia, mes, año.
    return day, month, year

# -----------------------------------------------------------------------------



# Funcion para extraer el contenido de la tabla de excel.
# -----------------------------------------------------------------------------
def read_excel(excelFile):
    # 1.Guarda el file.
    workbook = load_workbook(excelFile)
    
    # 2.Escoge el sheet.
    sheet = workbook["General Matrix"]

    # 3.Listas para guardar la data del file de excel.
    '''
        En esta parte va la informacion que quiere el usuario.
    '''
    caseNums = []
    nameList = []
    addrList = []
    ogpNums = []


    # 4.Guarda la data requerida de la tabla de excel en las listas.
    for row in range(3, sheet.max_row+1):
        # 1.Checkea si hay un valor en la tabla de excel.
        if sheet.cell(row, 3).value: #Si ha un valor el la fila x, columna 3, entonces
            # 2. guarda el valor de la data requerida en las listas.
            caseNums.append(sheet.cell(row, 2).value)
            nameList.append(sheet.cell(row, 3).value)
            ogpNums.append(sheet.cell(row, 26).value)
            # Combina los valor de la tabla de excel para hacer la direccion.
            address = str(sheet.cell(row, 4).value) + ", " + str(sheet.cell(row, 5).value) + str(sheet.cell(row, 6).value)
            addrList.append(address)

    # Guarda las 3 listas.
    return caseNums, nameList, addrList, ogpNums
# -----------------------------------------------------------------------------

# Guarda la fecha de creacion de formas.
day, month, year = get_date()


# Funcion para editar la forma1.
'''
    contexto: Esta funcion corre en un for loop para crear una forma por cada cliente
                de la tabla de excel.
    parametros: 
        1.directorio - nombre del folder donde se van a guardar los files.
        2.word - file que sirve como base para las formas que se van a crear.
        3.name - nombre de cliente.
        4.number - numero del cliente.
        5.address - direccion del cliente.
    
'''
# -----------------------------------------------------------------------------
def form_one(directory, word, name, number, address):
    # 1.Escribe la fecha en el formato deseado.
    fecha = day + ' de ' + month + ' de ' + year
    # 2.Guarda la espacio en la memoria para la forma 1. 
    wordDoc = Document(word)
    
    
    # 3.por cada texto, encuentra el texto marcado y lo remplaza con la data deseada.
    for paragraph in wordDoc.paragraphs:
        # Si el parrafo tiene :KEYWORD: entonces remplazalo con la :Variable:
        if '#CASE#' in paragraph.text:
            paragraph.text = paragraph.text.replace("#CASE#", str(number))
        if '#NAME#' in paragraph.text:
            paragraph.text = paragraph.text.replace("#NAME#", str(name))
        if '#ADDR#' in paragraph.text:
            paragraph.text = paragraph.text.replace("#ADDR#", str(address))
        if '#DATE#' in paragraph.text:
            paragraph.text = paragraph.text.replace("#DATE#", str(fecha))

    # 4.Crea el path para las formas creadas.
    newName = directory + "/" + word.replace(".docx", "_" + name.replace(" ", "-") + '-' + number + ".docx")
    # (4)Guarda el documento en el path creado.
    wordDoc.save(newName)
# -----------------------------------------------------------------------------


# Funcion para editar la forma2.
'''
    contexto: Esta funcion corre en un for loop para crear una forma por cada cliente
                de la tabla de excel.
    parametros: 
        1.directorio - nombre del folder donde se van a guardar los files.
        2.word - file que sirve como base para las formas que se van a crear.
        3.name - nombre de cliente.
        4.number - numero del cliente.
        5.address - direccion del cliente.
    
'''
# -----------------------------------------------------------------------------
def form_two(directory, word, name, number, ogpnum):
    # 1.Guarda la forma 2.
    wordDoc = Document(word)
    

    
    # 3.por cada texto, encuentra el texto marcado y lo remplaza con la data deseada.
    for table in wordDoc.tables:
        #  Itera sobre las filas de las tablas
        for row in table.rows:
            # Itera sobre los valores de cada fila.
            for cell in row.cells:
                # Itera sobre el texto de cada valor.
                for paragraph in cell.paragraphs:

                     # Si el texto tiene el :KEYWORD: entonces reemplazalo con la :Variable:
                    if '#PNAME#' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#PNAME#", str(number) + ' ' + str(name))
                    if '#OGPNUM#' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#OGPNUM#", str(ogpnum))
                    if '#D' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#D", day)
                    if '#M' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#M", month)
                    if '#Y' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#Y", year)


    # 4.Crea el path para las formas creadas.
    newName = directory + "/" + word.replace(".docx", "_" + name.replace(" ", "-") + '-' + number + ".docx")
    # (4)Guarda el documento en el path creado.
    wordDoc.save(newName)
    

# Funcion para editar la forma 3.
def form_three(directory, word, name, number, ogpnum):
    # Guarda la base de la forma 3.
    wordDoc = Document(word)
    # Crea un path nuevo para la forma 3.
    newName = directory + "/" + word.replace(".docx", "_" + name.replace(" ", "-") + '-' + number + ".docx")
    

    
    # 3.por cada texto, encuentra el texto marcado y lo remplaza con la data deseada.
    for table in wordDoc.tables:
        #  Itera sobre las filas de las tablas
        for row in table.rows:
            # Itera sobre los valores de cada fila.
            for cell in row.cells:
                # Itera sobre el texto de cada valor.
                for paragraph in cell.paragraphs:

                    # Si el texto tiene el :KEYWORD: entonces reemplazalo con la :Variable:
                    if '#PNAME#' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#PNAME#", str(number) + ' ' + str(name))
                    if '#OGPNUM#' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#OGPNUM#", str(ogpnum))
                    if '#D' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#D", day)
                    if '#M' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#M", month)
                    if '#Y' in paragraph.text:
                        paragraph.text = paragraph.text.replace("#Y", year)

    # Guarda la forma creada en el path creado.
    wordDoc.save(newName)
    
# -----------------------------------------------------------------------------


# -----------------------------------------------------------------------------
# DRIVER CODE:
#Nombre de la tabla de excel.
xlsxname = 'TCG Cases Services Requisition (ANDREA).xlsx'
caseNumbers, names, addresses, ogpNumbers = read_excel(xlsxname)

# Nombre del folder para guardar las formas creadas.
dirName = "./Word Files"
# Si el folder no existe crearlo, si ya existe, continua.
try:
    os.mkdir(dirName)
except Exception:
    pass

# For loop para ejecutar las funciones que crean las formas con la informacion de cada cliente.
for i in range(len(names)):
    form_three(dirName, 'form20.docx', names[i], caseNumbers[i], ogpNumbers[i])
    form_two(dirName, 'form11.docx', names[i], caseNumbers[i], ogpNumbers[i])
    form_one(dirName, "form5.docx", names[i], caseNumbers[i], addresses[i])
    
#Crea y guarda los doc files en formato pdf.
convert('Word Files/')
# -----------------------------------------------------------------------------
